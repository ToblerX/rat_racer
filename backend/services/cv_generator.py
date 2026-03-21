import io
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

URL_PATTERN = re.compile(r"(https?://[^\s\)]+)")
LINK_WITH_LABEL_PATTERN = re.compile(r"([^(\n]+?)\s*\((https?://[^\s\)]+)\)")
BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")

# Combined pattern to split text into: links, bold markers, and plain text
# Order matters: check links first, then bold, then bare URLs
INLINE_PATTERN = re.compile(
    r"([^(\n]+?)\s*\((https?://[^\s\)]+)\)"  # Label (URL)
    r"|\*\*(.+?)\*\*"  # **bold**
    r"|(https?://[^\s\)]+)"  # bare URL
)


def add_hyperlink(paragraph, text, url, font_size=None):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)

    sz_val = str(font_size * 2) if font_size else "20"
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), sz_val)
    rPr.append(sz)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)

    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def add_rich_text(paragraph, text, base_font_size=10):
    """Add text to a paragraph with hyperlinks (label only) and **bold** support."""
    last_end = 0
    for match in INLINE_PATTERN.finditer(text):
        # Add plain text before this match
        before = text[last_end : match.start()]
        if before:
            run = paragraph.add_run(before)
            run.font.size = Pt(base_font_size)
            run.font.name = "Calibri"

        if match.group(2):
            # Label (URL) — render label as clickable hyperlink
            label = match.group(1).strip()
            url = match.group(2).strip()
            add_hyperlink(paragraph, label, url, font_size=base_font_size)
        elif match.group(3):
            # **bold text**
            run = paragraph.add_run(match.group(3))
            run.bold = True
            run.font.size = Pt(base_font_size)
            run.font.name = "Calibri"
        elif match.group(4):
            # Bare URL — render as clickable
            url = match.group(4).strip()
            add_hyperlink(paragraph, url, url, font_size=base_font_size)

        last_end = match.end()

    # Remaining text after last match
    remaining = text[last_end:]
    if remaining:
        run = paragraph.add_run(remaining)
        run.font.size = Pt(base_font_size)
        run.font.name = "Calibri"


SECTION_HEADERS = {
    "summary",
    "skills",
    "experience",
    "education",
    "languages",
    "projects",
    "certifications",
    "interests",
    "references",
    "professional experience",
    "work experience",
    "technical skills",
}

# Pattern for contact info lines: phone | email | location
CONTACT_PATTERN = re.compile(
    r"[\+\d\s\-\(\)]{7,}|[\w\.\-]+@[\w\.\-]+\.\w+|[A-Z][a-z]+,\s*[A-Z]"
)

# Pattern for links line (e.g. "Website GitHub LinkedIn")
LINKS_LINE_PATTERN = re.compile(
    r"^(Website|GitHub|LinkedIn|Instagram|Portfolio|Blog|Twitter|X)(\s+(Website|GitHub|LinkedIn|Instagram|Portfolio|Blog|Twitter|X))*$",
    re.IGNORECASE,
)


def add_separator(doc):
    """Add a thin horizontal line separator."""
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing."""
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def is_section_header(text: str) -> bool:
    """Check if a line is a CV section header."""
    clean = text.strip().rstrip(":").lower()
    return clean in SECTION_HEADERS or text.strip().isupper()


def is_job_title_line(text: str) -> bool:
    """Detect job title / date lines like 'Fullstack Engineer  December 2025 – Now'."""
    return bool(
        re.search(
            r"(January|February|March|April|May|June|July|August|September|October|November|December|\d{4})\s*[–\-—]\s*(Now|Present|Current|January|February|March|April|May|June|July|August|September|October|November|December|\d{4})",
            text,
            re.IGNORECASE,
        )
    )


def is_company_line(text: str) -> bool:
    """Detect company/location lines like 'Teachapp (solo, self-employed)  Remote'."""
    return bool(re.search(r"Remote$|Hybrid$|On-?site$", text.strip(), re.IGNORECASE))


def generate_docx(cv_text: str) -> bytes:
    """Generate a professionally styled DOCX file from plain CV text."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Compact default paragraph spacing
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    lines = cv_text.split("\n")
    is_first_line = True
    name_written = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        # --- Name: first non-empty line, centered, uppercase, large ---
        if is_first_line:
            is_first_line = False
            p = doc.add_paragraph()
            p.alignment = 1  # CENTER
            run = p.add_run(stripped.upper())
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            run.font.name = "Calibri"
            set_paragraph_spacing(p, before=0, after=40)
            name_written = True
            continue

        # --- Contact info line (phone | email | location) ---
        if name_written and CONTACT_PATTERN.search(stripped) and "|" in stripped:
            p = doc.add_paragraph()
            p.alignment = 1  # CENTER
            add_rich_text(p, stripped, base_font_size=9)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            set_paragraph_spacing(p, before=0, after=20)
            continue

        # --- Links line (Website GitHub LinkedIn etc.) or lines with Label (URL) patterns in header area ---
        if name_written and (
            LINKS_LINE_PATTERN.match(stripped)
            or LINK_WITH_LABEL_PATTERN.search(stripped)
        ):
            p = doc.add_paragraph()
            p.alignment = 1  # CENTER
            add_rich_text(p, stripped, base_font_size=9)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            set_paragraph_spacing(p, before=0, after=20)
            # Check if next non-empty line is a section header to end header block
            for j in range(i + 1, len(lines)):
                next_stripped = lines[j].strip()
                if next_stripped:
                    if is_section_header(next_stripped):
                        name_written = False
                    break
            continue

        # After contact/links block, stop special header handling
        if name_written and not CONTACT_PATTERN.search(stripped):
            name_written = False

        # --- Section headers (SUMMARY, SKILLS, EXPERIENCE, etc.) ---
        if is_section_header(stripped):
            add_separator(doc)
            p = doc.add_paragraph()
            run = p.add_run(stripped.rstrip(":").upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            run.font.name = "Calibri"
            run.font.all_caps = True
            set_paragraph_spacing(p, before=120, after=40)
            continue

        # --- Job title lines (with date range) ---
        if is_job_title_line(stripped):
            p = doc.add_paragraph()
            # Split into title and date parts
            parts = re.split(r"\t+|\s{2,}", stripped, maxsplit=1)
            run = p.add_run(parts[0])
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            run.font.name = "Calibri"
            if len(parts) > 1:
                date_run = p.add_run("  |  " + parts[1])
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                date_run.font.name = "Calibri"
            set_paragraph_spacing(p, before=80, after=0)
            continue

        # --- Company / location lines ---
        if is_company_line(stripped):
            p = doc.add_paragraph()
            parts = re.split(r"\t+|\s{2,}", stripped, maxsplit=1)
            run = p.add_run(parts[0])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = "Calibri"
            if len(parts) > 1:
                loc_run = p.add_run("  |  " + parts[1])
                loc_run.italic = True
                loc_run.font.size = Pt(9)
                loc_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                loc_run.font.name = "Calibri"
            set_paragraph_spacing(p, before=0, after=20)
            continue

        # --- Bullet points ---
        if stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, stripped[2:], base_font_size=10)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            set_paragraph_spacing(p, before=0, after=10)
            continue

        # --- Skills lines with tab/colon separators (e.g. "Languages: Python, TypeScript") ---
        if (
            ":" in stripped
            and len(stripped.split(":")[0]) < 30
            and "http" not in stripped.split(":")[0]
        ):
            label, _, value = stripped.partition(":")
            p = doc.add_paragraph()
            label_run = p.add_run(label.strip() + ": ")
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            label_run.font.name = "Calibri"
            add_rich_text(p, value.strip(), base_font_size=10)
            for run in p.runs:
                if not run.bold:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            set_paragraph_spacing(p, before=0, after=10)
            continue

        # --- Default paragraph ---
        p = doc.add_paragraph()
        add_rich_text(p, stripped, base_font_size=10)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        set_paragraph_spacing(p, before=0, after=20)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
